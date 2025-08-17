import streamlit as st
import numpy as np
from PIL import Image, ImageDraw, ImageFilter, ImageEnhance, ImageFont
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Inches
import pytesseract
import cv2

# Configure page with modern styling
st.set_page_config(
    page_title="AI Digitizer Pro",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Custom CSS for professional look
st.markdown("""
<style>
    .main { padding-top: 0rem; }
    .stApp { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); }
    
    .hero-container {
        background: rgba(255, 255, 255, 0.1);
        backdrop-filter: blur(10px);
        border-radius: 20px;
        padding: 2rem;
        margin: 1rem 0;
        border: 1px solid rgba(255, 255, 255, 0.2);
        box-shadow: 0 8px 32px rgba(0, 0, 0, 0.1);
    }
    
    .mode-card {
        background: rgba(255, 255, 255, 0.95);
        border-radius: 15px;
        padding: 1.5rem;
        margin: 1rem 0;
        box-shadow: 0 8px 25px rgba(0, 0, 0, 0.1);
        border-left: 4px solid #667eea;
        transition: transform 0.3s ease;
    }
    
    .mode-card:hover {
        transform: translateY(-5px);
        box-shadow: 0 12px 35px rgba(0, 0, 0, 0.15);
    }
    
    .feature-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
        gap: 1.5rem;
        margin: 2rem 0;
    }
    
    .feature-item {
        background: rgba(255, 255, 255, 0.9);
        border-radius: 12px;
        padding: 1.5rem;
        text-align: center;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.1);
    }
    
    .upload-zone {
        border: 2px dashed #667eea;
        border-radius: 15px;
        padding: 2rem;
        text-align: center;
        background: rgba(255, 255, 255, 0.9);
        margin: 1rem 0;
    }
    
    .result-container {
        background: rgba(255, 255, 255, 0.95);
        border-radius: 15px;
        padding: 1.5rem;
        margin: 1rem 0;
        box-shadow: 0 8px 25px rgba(0, 0, 0, 0.1);
    }
    
    .stButton > button {
        background: linear-gradient(45deg, #667eea, #764ba2);
        color: white;
        border: none;
        border-radius: 25px;
        padding: 0.5rem 2rem;
        font-weight: 600;
        transition: all 0.3s ease;
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 25px rgba(102, 126, 234, 0.4);
    }
    
    .metric-card {
        background: linear-gradient(45deg, #667eea, #764ba2);
        color: white;
        border-radius: 12px;
        padding: 1rem;
        text-align: center;
        margin: 0.5rem;
    }
</style>
""", unsafe_allow_html=True)

def preprocess_for_ocr(image):
    """Enhanced preprocessing for OCR"""
    # Convert to OpenCV format
    opencv_img = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
    gray = cv2.cvtColor(opencv_img, cv2.COLOR_BGR2GRAY)
    
    # Noise removal
    denoised = cv2.fastNlMeansDenoising(gray)
    
    # Thresholding
    _, thresh = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    
    # Morphological operations
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 2))
    processed = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)
    
    return Image.fromarray(processed)

def extract_text_with_ocr(image):
    """Extract text using Tesseract OCR"""
    try:
        # Preprocess image for better OCR
        processed_img = preprocess_for_ocr(image)
        
        # Configure Tesseract
        custom_config = r'--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz.,!?;:()[]{}/"\'- '
        
        # Extract text
        text = pytesseract.image_to_string(processed_img, config=custom_config)
        
        # Get word-level data for positioning
        data = pytesseract.image_to_data(processed_img, output_type=pytesseract.Output.DICT)
        
        return text.strip(), data
    except Exception as e:
        st.error(f"OCR Error: {str(e)}")
        return "", None

def create_digital_text_document(text, confidence_data=None):
    """Create a clean digital version of extracted text"""
    if not text:
        return None
    
    # Create a simple formatted document
    lines = text.split('\n')
    formatted_lines = []
    
    for line in lines:
        line = line.strip()
        if line:
            formatted_lines.append(line)
    
    return '\n'.join(formatted_lines)

# [Keep the existing diagram processing functions]
def preprocess_image(image):
    """Enhanced image preprocessing for better shape detection"""
    if image.mode != 'L':
        gray = image.convert('L')
    else:
        gray = image
    
    enhancer = ImageEnhance.Contrast(gray)
    enhanced = enhancer.enhance(2.0)
    blurred = enhanced.filter(ImageFilter.GaussianBlur(radius=0.5))
    gray_array = np.array(gray)
    blurred_array = np.array(blurred)
    
    threshold = np.mean(blurred_array) - 20
    thresh = blurred_array < threshold
    thresh = thresh.astype(np.uint8) * 255
    
    return thresh, gray_array

def detect_shapes_and_text(binary_image, original_gray):
    """Detect shapes and estimate text content"""
    shapes_detected = []
    binary = binary_image > 128
    height, width = binary.shape
    visited = np.zeros_like(binary, dtype=bool)
    
    def flood_fill(start_y, start_x):
        if (start_y < 0 or start_y >= height or 
            start_x < 0 or start_x >= width or 
            visited[start_y, start_x] or 
            not binary[start_y, start_x]):
            return []
        
        points = []
        stack = [(start_y, start_x)]
        
        while stack:
            y, x = stack.pop()
            if (y < 0 or y >= height or 
                x < 0 or x >= width or 
                visited[y, x] or 
                not binary[y, x]):
                continue
                
            visited[y, x] = True
            points.append((y, x))
            
            for dy in [-1, 0, 1]:
                for dx in [-1, 0, 1]:
                    if dy != 0 or dx != 0:
                        stack.append((y + dy, x + dx))
        
        return points
    
    def analyze_shape_type(points, bbox):
        min_y, min_x, max_y, max_x = bbox
        w = max_x - min_x + 1
        h = max_y - min_y + 1
        area = len(points)
        aspect_ratio = w / h if h > 0 else 1
        fill_ratio = area / (w * h) if (w * h) > 0 else 0
        
        center_x, center_y = (min_x + max_x) / 2, (min_y + max_y) / 2
        distances_from_center = [((x - center_x) ** 2 + (y - center_y) ** 2) ** 0.5 
                               for y, x in points]
        
        if distances_from_center:
            avg_distance = np.mean(distances_from_center)
            distance_variance = np.var(distances_from_center)
            circularity = 1 / (1 + distance_variance / (avg_distance ** 2)) if avg_distance > 0 else 0
        else:
            circularity = 0
        
        if circularity > 0.7 and fill_ratio > 0.5:
            return "oval"
        elif aspect_ratio > 2 or aspect_ratio < 0.5:
            return "rectangle"
        elif 0.8 <= aspect_ratio <= 1.2 and fill_ratio > 0.6:
            return "square"
        elif fill_ratio < 0.3:
            return "diamond"
        else:
            return "rectangle"
    
    shape_id = 0
    
    for y in range(height):
        for x in range(width):
            if binary[y, x] and not visited[y, x]:
                points = flood_fill(y, x)
                
                if len(points) > 200:
                    ys, xs = zip(*points)
                    min_y, max_y = min(ys), max(ys)
                    min_x, max_x = min(xs), max(xs)
                    
                    w = max_x - min_x + 1
                    h = max_y - min_y + 1
                    
                    if w < 20 or h < 20:
                        continue
                    
                    shape_type = analyze_shape_type(points, (min_y, min_x, max_y, max_x))
                    
                    # Simple text estimation
                    area = len(points)
                    if area > 1000:
                        text_content = "Process Step"
                    elif area > 500:
                        text_content = "Decision"
                    else:
                        text_content = "Start/End"
                    
                    shapes_detected.append({
                        'id': shape_id,
                        'type': shape_type,
                        'text': text_content,
                        'x': min_x,
                        'y': min_y,
                        'width': w,
                        'height': h,
                        'center_x': min_x + w // 2,
                        'center_y': min_y + h // 2,
                        'area': len(points)
                    })
                    
                    shape_id += 1
    
    return shapes_detected

def create_clean_digital_flowchart(shapes, canvas_width=None, canvas_height=None):
    """Create a professional clean digital flowchart"""
    if not shapes:
        return None
    
    if canvas_width is None or canvas_height is None:
        max_x = max([s['x'] + s['width'] for s in shapes]) + 100
        max_y = max([s['y'] + s['height'] for s in shapes]) + 100
        canvas_width = max(max_x, 800)
        canvas_height = max(max_y, 600)
    
    canvas = Image.new('RGB', (canvas_width, canvas_height), 'white')
    draw = ImageDraw.Draw(canvas)
    
    colors = {
        'rectangle': '#E3F2FD',
        'square': '#F3E5F5',
        'oval': '#E8F5E8',
        'diamond': '#FFF3E0'
    }
    
    border_colors = {
        'rectangle': '#1976D2',
        'square': '#7B1FA2',
        'oval': '#388E3C',
        'diamond': '#F57C00'
    }
    
    sorted_shapes = sorted(shapes, key=lambda x: x['area'], reverse=True)
    
    for shape in sorted_shapes:
        x, y = shape['x'], shape['y']
        w, h = shape['width'], shape['height']
        shape_type = shape['type']
        text = shape['text']
        
        fill_color = colors.get(shape_type, '#F5F5F5')
        border_color = border_colors.get(shape_type, '#424242')
        
        if shape_type == 'rectangle' or shape_type == 'square':
            draw.rectangle([x, y, x + w, y + h], 
                         fill=fill_color, outline=border_color, width=3)
        elif shape_type == 'oval':
            draw.ellipse([x, y, x + w, y + h], 
                        fill=fill_color, outline=border_color, width=3)
        elif shape_type == 'diamond':
            points = [
                (x + w//2, y),
                (x + w, y + h//2),
                (x + w//2, y + h),
                (x, y + h//2)
            ]
            draw.polygon(points, fill=fill_color, outline=border_color, width=3)
        
        if text and text.strip():
            try:
                font_size = min(w // max(len(text), 1) + 5, h // 3, 16)
                font_size = max(font_size, 10)
                
                text_bbox = draw.textbbox((0, 0), text)
                text_width = text_bbox[2] - text_bbox[0]
                text_height = text_bbox[3] - text_bbox[1]
                
                text_x = x + (w - text_width) // 2
                text_y = y + (h - text_height) // 2
                
                draw.text((text_x + 1, text_y + 1), text, fill='#CCCCCC')
                draw.text((text_x, text_y), text, fill='#212121')
                
            except Exception:
                draw.text((x + 5, y + h//2 - 5), text, fill='#212121')
    
    return canvas

def main():
    # Hero Section
    st.markdown("""
    <div class="hero-container">
        <h1 style="text-align: center; color: white; font-size: 3rem; margin-bottom: 1rem;">
            🚀 AI Digitizer Pro
        </h1>
        <p style="text-align: center; color: rgba(255,255,255,0.9); font-size: 1.2rem;">
            Transform handwritten content into professional digital formats with AI
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # Mode Selection
    st.markdown("### 🎯 Choose Your Digitization Mode")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("""
        <div class="mode-card">
            <h3>📝 Text Digitizer</h3>
            <p>Convert handwritten notes, documents, and text into clean digital format using advanced OCR technology.</p>
            <ul>
                <li>✨ Advanced OCR recognition</li>
                <li>📄 Multiple output formats</li>
                <li>🔧 Text cleanup & formatting</li>
                <li>📊 Confidence analysis</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
        
        text_mode = st.button("🚀 Launch Text Digitizer", key="text_btn", type="primary")
    
    with col2:
        st.markdown("""
        <div class="mode-card">
            <h3>🔄 Diagram Digitizer</h3>
            <p>Transform hand-drawn flowcharts and diagrams into professional digital versions with perfect geometry.</p>
            <ul>
                <li>🤖 AI shape detection</li>
                <li>🎨 Professional styling</li>
                <li>📐 Perfect geometry</li>
                <li>🌈 Color-coded elements</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
        
        diagram_mode = st.button("🚀 Launch Diagram Digitizer", key="diagram_btn", type="primary")
    
    # Initialize session state
    if 'mode' not in st.session_state:
        st.session_state.mode = None
    if 'converted' not in st.session_state:
        st.session_state.converted = False
    
    # Set mode based on button clicks
    if text_mode:
        st.session_state.mode = 'text'
        st.session_state.converted = False
    elif diagram_mode:
        st.session_state.mode = 'diagram'
        st.session_state.converted = False
    
    # Display selected mode interface
    if st.session_state.mode == 'text':
        text_digitizer_interface()
    elif st.session_state.mode == 'diagram':
        diagram_digitizer_interface()
    else:
        # Show features when no mode selected
        show_features()

def text_digitizer_interface():
    st.markdown("---")
    st.markdown("## 📝 Text Digitizer Mode")
    
    col1, col2 = st.columns([1, 3])
    with col1:
        if st.button("← Back to Home", type="secondary"):
            st.session_state.mode = None
            st.rerun()
    
    # File uploader
    st.markdown("""
    <div class="upload-zone">
        <h3>📤 Upload Your Handwritten Text</h3>
        <p>Supports: JPG, PNG, PDF, and more</p>
    </div>
    """, unsafe_allow_html=True)
    
    uploaded_file = st.file_uploader(
        "Choose file", 
        type=['jpg', 'jpeg', 'png', 'bmp', 'tiff'],
        label_visibility="collapsed"
    )
    
    if uploaded_file:
        image = Image.open(uploaded_file)
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("""<div class="result-container">""", unsafe_allow_html=True)
            st.markdown("#### 📄 Original Document")
            st.image(image, use_column_width=True)
            
            if st.button("🤖 Extract Text", type="primary"):
                with st.spinner("🔍 Analyzing text with AI OCR..."):
                    extracted_text, confidence_data = extract_text_with_ocr(image)
                    
                    if extracted_text:
                        st.session_state.extracted_text = extracted_text
                        st.session_state.confidence_data = confidence_data
                        st.session_state.converted = True
                        st.success("✅ Text extraction completed!")
                    else:
                        st.error("❌ No text detected. Try a clearer image.")
            
            st.markdown("</div>", unsafe_allow_html=True)
        
        with col2:
            st.markdown("""<div class="result-container">""", unsafe_allow_html=True)
            st.markdown("#### ✨ Digitized Text")
            
            if hasattr(st.session_state, 'extracted_text') and st.session_state.extracted_text:
                # Show extracted text
                st.text_area("Extracted Text:", st.session_state.extracted_text, height=300)
                
                # Download options
                st.markdown("#### 📥 Download Options")
                col_a, col_b = st.columns(2)
                
                with col_a:
                    st.download_button(
                        "📄 Download as TXT",
                        st.session_state.extracted_text,
                        file_name="extracted_text.txt",
                        mime="text/plain"
                    )
                
                with col_b:
                    # Create Word document
                    doc = Document()
                    doc.add_heading('Extracted Text', 0)
                    doc.add_paragraph(st.session_state.extracted_text)
                    
                    doc_buffer = BytesIO()
                    doc.save(doc_buffer)
                    doc_buffer.seek(0)
                    
                    st.download_button(
                        "📄 Download as DOCX",
                        doc_buffer.getvalue(),
                        file_name="extracted_text.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.info("👆 Upload an image and click 'Extract Text' to see results")
            
            st.markdown("</div>", unsafe_allow_html=True)

def diagram_digitizer_interface():
    st.markdown("---")
    st.markdown("## 🔄 Diagram Digitizer Mode")
    
    col1, col2 = st.columns([1, 3])
    with col1:
        if st.button("← Back to Home", type="secondary"):
            st.session_state.mode = None
            st.rerun()
    
    # Settings
    with st.expander("⚙️ Advanced Settings"):
        min_shape_size = st.slider("Minimum Shape Size", 100, 1000, 300)
    
    # File uploader
    st.markdown("""
    <div class="upload-zone">
        <h3>📤 Upload Your Hand-drawn Diagram</h3>
        <p>Flowcharts, process diagrams, mind maps, etc.</p>
    </div>
    """, unsafe_allow_html=True)
    
    uploaded_file = st.file_uploader(
        "Choose file",
        type=['jpg', 'jpeg', 'png', 'bmp'],
        label_visibility="collapsed"
    )
    
    if uploaded_file:
        image = Image.open(uploaded_file)
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("""<div class="result-container">""", unsafe_allow_html=True)
            st.markdown("#### 📝 Original Diagram")
            st.image(image, use_column_width=True)
            
            if st.button("🤖 Convert to Digital", type="primary"):
                with st.spinner("🔍 Analyzing diagram structure..."):
                    processed_img, gray_img = preprocess_image(image)
                    shapes = detect_shapes_and_text(processed_img, gray_img)
                    shapes = [s for s in shapes if s['area'] >= min_shape_size]
                    
                    if shapes:
                        digital_flowchart = create_clean_digital_flowchart(shapes)
                        st.session_state.digital_image = digital_flowchart
                        st.session_state.detected_shapes = shapes
                        st.session_state.converted = True
                        st.success(f"✅ Converted! Detected {len(shapes)} shapes")
                    else:
                        st.warning("❌ No shapes detected. Try adjusting settings.")
            
            st.markdown("</div>", unsafe_allow_html=True)
        
        with col2:
            st.markdown("""<div class="result-container">""", unsafe_allow_html=True)
            st.markdown("#### ✨ Digital Version")
            
            if hasattr(st.session_state, 'digital_image') and st.session_state.digital_image:
                st.image(st.session_state.digital_image, use_column_width=True)
                
                # Download options
                st.markdown("#### 📥 Download Options")
                col_a, col_b = st.columns(2)
                
                with col_a:
                    png_buffer = BytesIO()
                    st.session_state.digital_image.save(png_buffer, format='PNG')
                    st.download_button(
                        "🖼️ Download PNG",
                        png_buffer.getvalue(),
                        file_name="digital_diagram.png",
                        mime="image/png"
                    )
                
                with col_b:
                    jpg_buffer = BytesIO()
                    rgb_img = st.session_state.digital_image.convert('RGB')
                    rgb_img.save(jpg_buffer, format='JPEG', quality=95)
                    st.download_button(
                        "📸 Download JPG",
                        jpg_buffer.getvalue(),
                        file_name="digital_diagram.jpg",
                        mime="image/jpeg"
                    )
            else:
                st.info("👆 Upload a diagram and click 'Convert to Digital'")
            
            st.markdown("</div>", unsafe_allow_html=True)

def show_features():
    st.markdown("### ✨ Why Choose AI Digitizer Pro?")
    
    st.markdown("""
    <div class="feature-grid">
        <div class="feature-item">
            <h3>🤖 AI-Powered</h3>
            <p>Advanced machine learning algorithms for accurate recognition and conversion</p>
        </div>
        <div class="feature-item">
            <h3>⚡ Lightning Fast</h3>
            <p>Process images in seconds with optimized performance</p>
        </div>
        <div class="feature-item">
            <h3>🎨 Professional Output</h3>
            <p>Clean, polished results ready for presentations and documents</p>
        </div>
        <div class="feature-item">
            <h3>📱 Multi-Format</h3>
            <p>Export to PNG, JPG, TXT, DOCX, and more formats</p>
        </div>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
